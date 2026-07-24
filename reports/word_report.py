from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from core.metrics import (
    kpis, monthly_summary, weekly_summary, stop_frequency, combinations,
    weekly_punctuality_by_route, weekly_route_usage,
)
from ai.conclusions import generate_conclusions
from ai.recommendations import generate_recommendations

def _add_table(doc, df, max_rows=20):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, column in enumerate(df.columns):
        table.rows[0].cells[i].text = str(column)
    for _, row in df.head(max_rows).iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if isinstance(value, float):
                cells[i].text = f"{value:.2f}"
            else:
                cells[i].text = str(value)
    return table

def build_word(df) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(.65)
    section.bottom_margin = Inches(.65)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INFORME EJECUTIVO\nCENTRO DE INTELIGENCIA OPERACIONAL KOA")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(18, 61, 118)
    period = f"Periodo: {df['fecha'].min():%d/%m/%Y} al {df['fecha'].max():%d/%m/%Y}"
    p = doc.add_paragraph(period)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Resumen ejecutivo", level=1)
    for item in generate_conclusions(df):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Indicadores principales", level=1)
    _add_table(doc, __import__("pandas").DataFrame([kpis(df)]))

    doc.add_heading("3. Análisis mensual", level=1)
    _add_table(doc, monthly_summary(df))

    doc.add_heading("4. Análisis semanal", level=1)
    _add_table(doc, weekly_summary(df), max_rows=30)

    doc.add_heading("5. Puntualidad semanal por recorrido", level=1)
    _add_table(doc, weekly_punctuality_by_route(df), max_rows=40)

    doc.add_heading("6. Uso semanal de las rutas", level=1)
    _add_table(doc, weekly_route_usage(df), max_rows=40)

    doc.add_heading("7. Frecuencia de paraderos", level=1)
    _add_table(doc, stop_frequency(df))

    doc.add_heading("8. Combinaciones y tiempos", level=1)
    _add_table(doc, combinations(df), max_rows=25)

    doc.add_heading("9. Plan de acción", level=1)
    _add_table(doc, __import__("pandas").DataFrame(generate_recommendations(df)))

    doc.add_heading("10. Nota metodológica", level=1)
    doc.add_paragraph(
        "La puntualidad oficial corresponde a salidas entre 0 y 5 minutos después de la hora programada. "
        "Los tiempos efectivos excluyen registros sin usuarios, sin paradas válidas, pendientes de validar y no ejecutados. "
        "La frecuencia de paraderos se obtiene directamente de la columna PARADAS."
    )

    output = BytesIO()
    doc.save(output)
    return output.getvalue()
